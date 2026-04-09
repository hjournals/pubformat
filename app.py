from flask import Flask, request, render_template_string, send_from_directory
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import uuid

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"

# ÖNEMLİ:
# template.dotx yerine Word'de açıp .docx olarak kaydetmen daha güvenlidir.
# Örneğin dosya adını "template.docx" yap.
TEMPLATE_FILE = "TEMPLATE.docx"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

HTML_FORM = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Makale Sistemi</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            max-width: 1000px;
            margin: 30px auto;
            line-height: 1.5;
            padding: 20px;
        }
        h1, h2 {
            margin-top: 24px;
            margin-bottom: 12px;
        }
        label {
            font-weight: bold;
            display: block;
            margin-top: 12px;
            margin-bottom: 6px;
        }
        input, textarea, select {
            width: 100%;
            padding: 10px;
            box-sizing: border-box;
            margin-bottom: 12px;
            font-size: 14px;
        }
        textarea {
            min-height: 120px;
        }
        .section {
            border: 1px solid #ddd;
            padding: 18px;
            margin-bottom: 20px;
            border-radius: 8px;
            background: #fafafa;
        }
        .hint {
            font-size: 13px;
            color: #555;
            margin-top: -6px;
            margin-bottom: 10px;
        }
        .checkbox-row {
            display: flex;
            align-items: center;
            gap: 8px;
            margin-top: 15px;
            margin-bottom: 20px;
        }
        .checkbox-row input {
            width: auto;
            margin: 0;
        }
        button {
            padding: 12px 20px;
            font-size: 16px;
            cursor: pointer;
        }
    </style>

    <script>
        function toggleLanguageFields() {
            const lang = document.getElementById("language").value;
            const trFields = document.querySelectorAll(".tr-field");
            const trInputs = document.querySelectorAll(".tr-input");

            if (lang === "tr") {
                trFields.forEach(el => el.style.display = "block");
                trInputs.forEach(el => el.required = true);
            } else {
                trFields.forEach(el => el.style.display = "none");
                trInputs.forEach(el => el.required = false);
            }
        }

        window.onload = function() {
            toggleLanguageFields();
        }
    </script>
</head>
<body>

    <h1>Makale Başvuru Formu</h1>

    <form method="POST" enctype="multipart/form-data">

        <div class="section">
            <h2>1. Makale Dili</h2>

            <label for="language">Makale Dili</label>
            <select name="language" id="language" onchange="toggleLanguageFields()" required>
                <option value="tr">Türkçe</option>
                <option value="en">English</option>
            </select>
        </div>

        <div class="section">
            <h2>2. Başlık Bilgileri</h2>

            <div class="tr-field">
                <label for="title_tr">Türkçe Başlık</label>
                <input type="text" name="title_tr" id="title_tr" class="tr-input">
            </div>

            <label for="title_en">English Title</label>
            <input type="text" name="title_en" id="title_en" required>
        </div>

        <div class="section">
            <h2>3. Özet Bilgileri</h2>

            <div class="tr-field">
                <label for="abstract_tr">Türkçe Özet</label>
                <textarea name="abstract_tr" id="abstract_tr" class="tr-input"></textarea>
            </div>

            <label for="abstract_en">Abstract</label>
            <textarea name="abstract_en" id="abstract_en" required></textarea>
        </div>

        <div class="section">
            <h2>4. Anahtar Kelimeler</h2>

            <div class="tr-field">
                <label for="keywords_tr">Türkçe Anahtar Kelimeler</label>
                <input type="text" name="keywords_tr" id="keywords_tr" class="tr-input">
            </div>

            <label for="keywords_en">Keywords</label>
            <input type="text" name="keywords_en" id="keywords_en" required>

            <div class="hint">
                Anahtar kelimeleri virgülle ayırarak yazın.
            </div>
        </div>

        <div class="section">
            <h2>5. Ana Metin Dosyası</h2>

            <label for="body_file">Makale Ana Metni (.docx)</label>
            <input type="file" name="body_file" id="body_file" accept=".docx" required>

            <div class="hint">
                Ana metni Word (.docx) dosyası olarak yükleyin.
            </div>
        </div>

        <div class="section">
            <h2>6. Ek Beyanlar</h2>

            <label for="ack">Acknowledgements</label>
            <textarea name="ack" id="ack"></textarea>

            <label for="funding">Funding</label>
            <textarea name="funding" id="funding"></textarea>

            <label for="conflict">Conflict of Interest</label>
            <textarea name="conflict" id="conflict"></textarea>
        </div>

        <div class="section">
            <h2>7. References</h2>

            <label for="references">References</label>
            <textarea name="references" id="references" required></textarea>

            <div class="hint">
                Tüm kaynakçayı toplu olarak yapıştırın. Her referansı ayrı satıra yazın.
                APA içeriğini yazar hazırlamalıdır; sistem yalnızca Word içinde düzen uygular.
            </div>
        </div>

        <div class="section">
            <h2>8. Kör Hakemlik</h2>

            <div class="checkbox-row">
                <input type="checkbox" name="blind" id="blind">
                <label for="blind" style="margin:0; font-weight:normal;">Kör hakem sürümü üret</label>
            </div>
        </div>

        <button type="submit">Word Oluştur</button>

    </form>

</body>
</html>
"""

RESULT_HTML = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Dosya Hazır</title>
</head>
<body style="font-family: Arial, sans-serif; max-width: 900px; margin: 30px auto;">
    <h2>Word dosyası hazır</h2>
    <p><a href="/download/{{ filename }}">Dosyayı indir</a></p>
    <p><a href="/">Yeni makale oluştur</a></p>
</body>
</html>
"""

def set_run_font(run, font_name="Times New Roman", font_size=None, bold=None, italic=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)

    if font_size is not None:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def format_article_title_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(12)
    fmt.line_spacing = 1

    if not paragraph.runs:
        run = paragraph.add_run(paragraph.text)
        set_run_font(run, font_name="Times New Roman", font_size=16, bold=True)
        return

    for run in paragraph.runs:
        set_run_font(run, font_name="Times New Roman", font_size=16, bold=True)

def fix_template_title(doc, article_title):
    """
    Yeni template'te başlık paragrafı Article Name / ArticleName stiliyle geliyor.
    Önce stile göre bulur, gerekirse metne göre eşleştirir.
    """
    normalized_title = (article_title or "").strip()

    # 1) Önce stile göre yakala
    for para in doc.paragraphs:
        style_name = ""
        try:
            if para.style:
                style_name = para.style.name or ""
        except Exception:
            pass

        if style_name.strip().lower() in ["article name", "articlename"]:
            format_article_title_paragraph(para)
            return True

    # 2) Stil adı yakalanmazsa metne göre bul
    if normalized_title:
        for para in doc.paragraphs:
            if (para.text or "").strip() == normalized_title:
                format_article_title_paragraph(para)
                return True

    return False

def append_body(target_doc, source_doc, skip_first_nonempty_paragraph=False):
    skipped = False

    for para in source_doc.paragraphs:
        para_text = (para.text or "").strip()

        if skip_first_nonempty_paragraph and not skipped:
            if para_text:
                skipped = True
                continue

        new_p = target_doc.add_paragraph()

        try:
            if para.style and para.style.name:
                new_p.style = para.style.name
        except Exception:
            pass

        for run in para.runs:
            new_run = new_p.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline

            font_name = run.font.name if run.font.name else "Times New Roman"
            font_size = run.font.size.pt if run.font.size else None
            set_run_font(new_run, font_name=font_name, font_size=font_size)

        new_p.alignment = para.alignment

        src_fmt = para.paragraph_format
        dst_fmt = new_p.paragraph_format
        dst_fmt.left_indent = src_fmt.left_indent
        dst_fmt.right_indent = src_fmt.right_indent
        dst_fmt.first_line_indent = src_fmt.first_line_indent
        dst_fmt.space_before = src_fmt.space_before
        dst_fmt.space_after = src_fmt.space_after
        dst_fmt.line_spacing = src_fmt.line_spacing
        dst_fmt.line_spacing_rule = src_fmt.line_spacing_rule

def clean_reference_lines(text):
    lines = [line.strip() for line in text.splitlines()]
    return [line for line in lines if line]

def add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(text)
    set_run_font(run, font_name="Times New Roman", font_size=12, bold=True)

    fmt = p.paragraph_format
    fmt.space_before = Pt(12)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1

def add_normal_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = p.add_run(text)
    set_run_font(run, font_name="Times New Roman", font_size=12)

    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1

    return p

def add_reference_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = p.add_run(text)
    set_run_font(run, font_name="Times New Roman", font_size=12)

    fmt = p.paragraph_format
    fmt.left_indent = Cm(1)
    fmt.first_line_indent = Cm(-1)
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1

    return p

def unique_filename(prefix="final", ext=".docx"):
    return f"{prefix}_{uuid.uuid4().hex}{ext}"

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        try:
            language = request.form.get("language", "").strip()

            title_tr = request.form.get("title_tr", "").strip()
            title_en = request.form.get("title_en", "").strip()

            abstract_tr = request.form.get("abstract_tr", "").strip()
            abstract_en = request.form.get("abstract_en", "").strip()

            keywords_tr = request.form.get("keywords_tr", "").strip()
            keywords_en = request.form.get("keywords_en", "").strip()

            ack = request.form.get("ack", "").strip()
            funding = request.form.get("funding", "").strip()
            conflict = request.form.get("conflict", "").strip()
            references = request.form.get("references", "").strip()

            if language not in ["tr", "en"]:
                return "Makale dili geçersiz."

            if language == "tr":
                if not title_tr or not abstract_tr or not keywords_tr:
                    return "Türkçe makalelerde Türkçe başlık, özet ve anahtar kelimeler zorunludur."
                if not title_en or not abstract_en or not keywords_en:
                    return "Türkçe makalelerde İngilizce başlık, abstract ve keywords de zorunludur."

            if language == "en":
                if not title_en or not abstract_en or not keywords_en:
                    return "İngilizce makalelerde English Title, Abstract ve Keywords zorunludur."

            author_block = "Anonymous Author(s)" if request.form.get("blind") else ""

            body_file = request.files.get("body_file")
            if not body_file or not body_file.filename.lower().endswith(".docx"):
                return "Lütfen body için .docx dosyası yükleyin."

            body_filename = unique_filename(prefix="body", ext=".docx")
            body_path = os.path.join(UPLOAD_FOLDER, body_filename)
            body_file.save(body_path)

            output_file = unique_filename(prefix="final", ext=".docx")
            output_path = os.path.join(OUTPUT_FOLDER, output_file)

            tpl = DocxTemplate(TEMPLATE_FILE)

            context = {
                "language": language,
                "title_tr": title_tr,
                "title_en": title_en,
                "author_block": author_block,
                "abstract_tr": abstract_tr,
                "abstract_en": abstract_en,
                "keywords_tr": keywords_tr,
                "keywords_en": keywords_en,
            }

            tpl.render(context)
            tpl.save(output_path)

            final_doc = Document(output_path)

            # Türkçe makalede ana başlık TR, İngilizce makalede EN
            article_title = title_tr if language == "tr" else title_en

            # Yeni template'teki Article Name stilini zorla düzelt
            fix_template_title(final_doc, article_title)

            final_doc.add_page_break()

            body_doc = Document(body_path)
            append_body(final_doc, body_doc, skip_first_nonempty_paragraph=True)

            if ack:
                final_doc.add_paragraph()
                add_heading(final_doc, "Acknowledgements")
                add_normal_paragraph(final_doc, ack)

            if funding:
                add_heading(final_doc, "Funding")
                add_normal_paragraph(final_doc, funding)

            if conflict:
                add_heading(final_doc, "Conflict of Interest")
                add_normal_paragraph(final_doc, conflict)

            add_heading(final_doc, "REFERENCES")
            for line in clean_reference_lines(references):
                add_reference_paragraph(final_doc, line)

            final_doc.save(output_path)

            return render_template_string(RESULT_HTML, filename=output_file)

        except Exception as e:
            return f"HATA: {str(e)}"

    return render_template_string(HTML_FORM)

@app.route("/download/<filename>")
def download(filename):
    return send_from_directory(OUTPUT_FOLDER, filename, as_attachment=True)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)
